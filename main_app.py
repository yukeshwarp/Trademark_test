# Version - 6.1  (Code Updated of Checking for Non-matching Class number)

from fileinput import filename
import time
import streamlit as st 
import pandas as pd
import fitz  # PyMuPDF
from pydantic import BaseModel, Field, ValidationError
from typing import List, Dict, Union
import base64
from docx import Document  
from docx.shared import Pt
from io import BytesIO
import re, ast

class TrademarkDetails(BaseModel):
    trademark_name: str = Field(description="The name of the Trademark", example="DISCOVER")
    status: str = Field(description="The Status of the Trademark", example="Registered")
    serial_number: str = Field(description="The Serial Number of the trademark from Chronology section", example="87−693,628")
    international_class_number: List[int] = Field(description="The International class number or Nice Classes number of the trademark from Goods/Services section or Nice Classes section", example=[18])
    owner: str = Field(description="The owner of the trademark", example="WALMART STORES INC")
    goods_services: str = Field(description="The goods/services from the document", example="LUGGAGE AND CARRYING BAGS; SUITCASES, TRUNKS, TRAVELLING BAGS, SLING BAGS FOR CARRYING INFANTS, SCHOOL BAGS; PURSES; WALLETS; RETAIL AND ONLINE RETAIL SERVICES")
    page_number: int = Field(description="The page number where the trademark details are found in the document", example=3)
    registration_number: Union[str, None] = Field(description="The Registration number of the trademark from Chronology section", example="5,809,957")
    design_phrase: str = Field(description="The design phrase of the trademark", example="THE MARK CONSISTS OF THE STYLIZED WORD 'MINI' FOLLOWED BY 'BY MOTHERHOOD.'", default="")
    

def preprocess_text(text: str) -> str:
    text = re.sub(r'\s+', ' ', text).strip()
    text = re.sub(r'[\u2013\u2014]', '-', text)
    return text

def is_correct_format_code1(page_text: str) -> bool:
    required_fields = ["Status:", "Goods/Services:"] # , "Last Reported Owner:"
    return all(field in page_text for field in required_fields)

def is_correct_format_code2(page_text: str) -> bool:
    required_fields = ["Register", "Nice Classes", "Goods & Services"]
    return all(field in page_text for field in required_fields)

def extract_trademark_details_code1(document_chunk: str) -> Dict[str, Union[str, List[int]]]:
    try:
        from openai import AzureOpenAI
        client = AzureOpenAI(  
                    azure_endpoint="https://theswedes.openai.azure.com/",  
                    api_key="783973291a7c4a74a1120133309860c0",  
                    api_version="2024-02-01",
                )  
        # messages=[
        #         {"role": "system", "content":"""You are an expert in analyzing Trademark Documents. Please carefully read the provided Trademark Documents above and extract the following key information:

        #                                         Extract these six (6) properties from the Trademark Document:
        #                                         - Trademark name: The name of the Trademark
        #                                         - Status: The Status of the Trademark
        #                                         - Serial number: The Serial Number of the trademark from Chronology section
        #                                         - International class number: The International class number or Nice Classes number of the trademark from Goods/Services section or Nice Classes section
        #                                         - Goods/Services: The goods/services from the document of the trademark with International class number
        #                                         - Owner: The owner of the trademark

        #                                         Guidelines:
        #                                         - The extracted information should be factual and accurate to the document.
        #                                         - Be extremely concise, except for the Goods/Services which should be copied in full.
        #                                         - The extracted entities should be self-contained and easily understood without the rest of the Trademark details.
        #                                         - If any property is missing from the Trademark, please try to re-extract the field once again, Yet property is missing from the Trademark, please leave the field empty rather than guessing.

        #                                         Extract the following details from the trademark document exactly: "trademark_name", "status", "serial_number", "international_class_number", "goods_services" and "owner". The "international_class_number" should be a list of dictionaries which contains one or more than one International class number.   
        #         """},
        #         {"role": "user", "content": 
        #             f"""  
        #             List of Trademark Documents:
        #             ---------------------------
        #             {document_chunk}
        #             ---------------------------
        #             """}
        #         ]
        
        messages=[
                {"role": "system", "content": "You are a helpful assistant for extracting Meta Data from the Trademark Document."},
                {"role": "user", "content": f"Extract the following details from the trademark document: trademark name, status, serial number, international class number, owner, goods & services, filed date, registration number.\n\nDocument:\n{document_chunk}"}
        ]
        response = client.chat.completions.create(  
                model="GPT35",  
                messages=messages,  
                temperature=0,  
                max_tokens=4000,  
        )  
        extracted_text = response.choices[0].message.content
        
        details = {}
        for line in extracted_text.split("\n"):
            if ":" in line:
                key, value = line.split(":", 1)
                details[key.strip().lower().replace(" ", "_")] = value.strip()
        return details
    
    except Exception as e:
        print(f"An error occurred: {e}")

def extract_registration_number(document: str) -> str:
    """ Extract the registration number from the Chronology section """
    match = re.search(r'Chronology:.*?Registration Number:\s*([\d,]+)', document, re.DOTALL)
    if match:
        return match.group(1).strip()
    return "No registration number presented in document"

def extract_trademark_details_code2(page_text: str) -> Dict[str, Union[str, List[int]]]:
    details = {}

    trademark_name_match = re.search(r"\d+\s*/\s*\d+\s*\n\s*\n\s*([A-Za-z0-9'&!,\-. ]+)\s*\n", page_text)
    if trademark_name_match:
        details["trademark_name"] = trademark_name_match.group(1).strip()
    else:
        trademark_name_match = re.search(r"(?<=\n)([A-Za-z0-9'&!,\-. ]+)(?=\n)", page_text)
        details["trademark_name"] = trademark_name_match.group(1).strip() if trademark_name_match else ""

    status_match = re.search(r'Status\s*(?:\n|:\s*)([A-Za-z]+)', page_text, re.IGNORECASE)
    details["status"] = status_match.group(1).strip() if status_match else ""

    owner_match = re.search(r'Holder\s*(?:\n|:\s*)(.*)', page_text, re.IGNORECASE)
    if owner_match:
        details["owner"] = owner_match.group(1).strip()
    else:
        owner_match = re.search(r'Owner\s*(?:\n|:\s*)(.*)', page_text, re.IGNORECASE)
        details["owner"] = owner_match.group(1).strip() if owner_match else ""
        
        

    nice_classes_match = re.search(r'Nice Classes\s*[\s:]*\n((?:\d+(?:,\s*\d+)*)\b)', page_text, re.IGNORECASE)
    if nice_classes_match:
        nice_classes_text = nice_classes_match.group(1)
        nice_classes = [int(cls.strip()) for cls in nice_classes_text.split(",")]
        details["international_class_number"] = nice_classes
    else:
        details["international_class_number"] = []
        


    serial_number_match = re.search(r'Application#\s*(.*)', page_text, re.IGNORECASE)
    details["serial_number"] = serial_number_match.group(1).strip() if serial_number_match else ""

    goods_services_match = re.search(r'Goods & Services\s*(.*?)(?=\s*G&S translation|$)', page_text, re.IGNORECASE | re.DOTALL)
    details["goods_services"] = goods_services_match.group(1).strip() if goods_services_match else ""
    
    registration_number_match = re.search(r'Registration#\s*(.*)', page_text, re.IGNORECASE)
    details["registration_number"] = registration_number_match.group(1).strip() if registration_number_match else ""
    
    # Description
    design_phrase = re.search(r'Description\s*(.*?)(?=\s*Applicant|Owner|Holder|$)', page_text, re.IGNORECASE | re.DOTALL)
    details["design_phrase"] = design_phrase.group(1).strip() if design_phrase else "No Design phrase presented in document"
    

    return details


def split_text(text: str, max_tokens: int = 1500) -> List[str]:
    chunks = []
    current_chunk = []
    current_length = 0

    for line in text.split('\n'):
        line_length = len(line.split())
        if current_length + line_length > max_tokens:
            chunks.append('\n'.join(current_chunk))
            current_chunk = [line]
            current_length = line_length
        else:
            current_chunk.append(line)
            current_length += line_length

    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    return chunks

def parse_international_class_numbers(class_numbers: str) -> List[int]:
    numbers = class_numbers.split(',')
    return [int(num.strip()) for num in numbers if num.strip().isdigit()]

# def extract_international_class_numbers_and_goods_services(document: str) -> Dict[str, Union[List[int], str]]:
#     """ Extract the International Class Numbers and Goods/Services from the document """
#     class_numbers = []
#     goods_services = []
#     pattern = r'International Class (\d+): (.*?)(?=\nInternational Class \d+:|\n[A-Z][a-z]+:|\nLast Reported Owner:|\Z)'
#     matches = re.findall(pattern, document, re.DOTALL)
#     for match in matches:
#         class_number = int(match[0])
#         class_numbers.append(class_number)
#         goods_services.append(f"Class {class_number}: {match[1].strip()}")
#     return {
#         "international_class_numbers": class_numbers,
#         "goods_services": "\n".join(goods_services)
#     }

def extract_international_class_numbers_and_goods_services(document: str, start_page: int, pdf_document: fitz.Document) -> Dict[str, Union[List[int], str]]:
    """ Extract the International Class Numbers and Goods/Services from the document over a range of pages """
    class_numbers = []
    goods_services = []
    combined_text = ""

    for i in range(start_page, min(start_page + 6, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_text += page_text
        if "Last Reported Owner:" in page_text:
            break

    pattern = r'International Class (\d+): (.*?)(?=\nInternational Class \d+:|\n[A-Z][a-z]+:|\nLast Reported Owner:|Disclaimers:|\Z)'
    matches = re.findall(pattern, combined_text, re.DOTALL)
    for match in matches:
        class_number = int(match[0])
        class_numbers.append(class_number)
        goods_services.append(f"Class {class_number}: {match[1].strip()}")

    return {
        "international_class_numbers": class_numbers,
        "goods_services": "\n".join(goods_services)
    }

def extract_design_phrase(document: str, start_page: int, pdf_document: fitz.Document) -> Dict[str, Union[List[int], str]]:
    """ Extract the design phrase from the document """
    combined_texts = ""
    for i in range(start_page, min(start_page + 8, pdf_document.page_count)):
        page = pdf_document.load_page(i)
        page_text = page.get_text()
        combined_texts += page_text
        if "Filing Correspondent:" in page_text:
            break
        
    pattern = r'Design Phrase:\s*(.*?)(?=Other U\.S\. Registrations:|Filing Correspondent:|Group:|USPTO Page:|$)'
    match = re.search(pattern, document, re.DOTALL) 
    if match:
        design_phrase = match.group(1).strip()
        # Remove any newline characters within the design phrase
        design_phrase = ' '.join(design_phrase.split())
        return design_phrase
    return "No Design phrase presented in document"

    
def parse_trademark_details(document_path: str) -> List[Dict[str, Union[str, List[int]]]]:
    with fitz.open(document_path) as pdf_document:
        all_extracted_data = []
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            page_text = page.get_text()
            
            if is_correct_format_code1(page_text):
                preprocessed_chunk = preprocess_text(page_text)
                extracted_data = extract_trademark_details_code1(preprocessed_chunk)
                additional_data = extract_international_class_numbers_and_goods_services(page_text, page_num, pdf_document)
                registration_number = extract_registration_number(page_text)
                design_phrase = extract_design_phrase(page_text, page_num, pdf_document)
                
                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    extracted_data.update(additional_data)
                    extracted_data["design_phrase"] = design_phrase
                    all_extracted_data.append(extracted_data)
                    extracted_data["registration_number"] = registration_number
                    
                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_name = data.get("trademark_name", "").split(',')[0].strip()
                        if "Global Filings" in trademark_name:
                            trademark_name = trademark_name.split("Global Filings")[0].strip()
                        owner = data.get("owner", "").split(',')[0].strip()
                        status = data.get("status", "").split(',')[0].strip()
                        serial_number = data.get("serial_number", "")
                        international_class_number = data.get("international_class_numbers", [])
                        goods_services = data.get("goods_services", "")
                        page_number = data.get("page_number", "")
                        registration_number = data.get("registration_number", "No registration number presented in document")
                        design_phrase = data.get("design_phrase", "No Design phrase presented in document")

                        # If crucial fields are missing, attempt to re-extract the values
                        if not trademark_name or not owner or not status or not international_class_number:
                            preprocessed_chunk = preprocess_text(data.get("raw_text", ""))
                            extracted_data = extract_trademark_details_code1(preprocessed_chunk)
                            trademark_name = extracted_data.get("trademark_name", trademark_name).split(',')[0].strip()
                            if "Global Filings" in trademark_name:
                                trademark_name = trademark_name.split("Global Filings")[0].strip()
                            owner = extracted_data.get("owner", owner).split(',')[0].strip()
                            status = extracted_data.get("status", status).split(',')[0].strip()
                            international_class_number = parse_international_class_numbers(extracted_data.get("international_class_number", "")) or international_class_number
                            registration_number = extracted_data.get("registration_number", registration_number).split(',')[0].strip()

                        trademark_details = TrademarkDetails(
                            trademark_name=trademark_name,
                            owner=owner,
                            status=status,
                            serial_number=serial_number,
                            international_class_number=international_class_number,
                            goods_services=goods_services,
                            page_number=page_number,
                            registration_number=registration_number,
                            design_phrase=design_phrase
                        )
                        trademark_info = {
                            "trademark_name": trademark_details.trademark_name,
                            "owner": trademark_details.owner,
                            "status": trademark_details.status,
                            "serial_number": trademark_details.serial_number,
                            "international_class_number": trademark_details.international_class_number,
                            "goods_services": trademark_details.goods_services,
                            "page_number": trademark_details.page_number,
                            "registration_number":trademark_details.registration_number,
                            "design_phrase": trademark_details.design_phrase
                        }
                        print(trademark_info)
                        print("_____________________________________________________________________________________________________________________________")
                        trademark_list.append(trademark_info)
                    except ValidationError as e:
                        print(f"Validation error for trademark {i}: {e}")
                                    
            else :
                if not is_correct_format_code2(page_text):
                    continue

                extracted_data = extract_trademark_details_code2(page_text)
                if extracted_data:
                    extracted_data["page_number"] = page_num + 1
                    all_extracted_data.append(extracted_data)

                trademark_list = []
                for i, data in enumerate(all_extracted_data, start=1):
                    try:
                        trademark_details = TrademarkDetails(
                            trademark_name=data.get("trademark_name", ""),
                            owner=data.get("owner", ""),
                            status=data.get("status", ""),
                            serial_number=data.get("serial_number", ""),
                            international_class_number=data.get("international_class_number", []),
                            goods_services=data.get("goods_services", ""),
                            page_number=data.get("page_number", 0),
                            registration_number=data.get("registration_number", ""),
                            design_phrase=data.get("design_phrase", "")
                        )
                        if (trademark_details.trademark_name != "" and trademark_details.owner != "" and trademark_details.status != "" and trademark_details.goods_services != ""):
                                trademark_info = {
                                    "trademark_name": trademark_details.trademark_name,
                                    "owner": trademark_details.owner,
                                    "status": trademark_details.status,
                                    "serial_number": trademark_details.serial_number,
                                    "international_class_number": trademark_details.international_class_number,
                                    "goods_services": trademark_details.goods_services,
                                    "page_number": trademark_details.page_number,
                                    "registration_number":trademark_details.registration_number,
                                    "design_phrase":trademark_details.design_phrase,
                                }
                                trademark_list.append(trademark_info)
                    except ValidationError as e:
                        print(f"Validation error for trademark {i}: {e}")

        return trademark_list

def compare_trademarks(existing_trademark: List[Dict[str, Union[str, List[int]]]], proposed_name: str, proposed_class: str, proposed_goods_services: str) -> List[Dict[str, int]]:
    proposed_classes = [int(c.strip()) for c in proposed_class.split(',')]
    messages=[
            {"role": "system", "content": """You are a trademark attorney to properly reasoning based on given conditions and assign conflict grade high or moderate or low to existing trademark and respond with only High or Moderate or Low. \n\n 
                                            Conditions for determining Conflict Grades:\n\n 
                                            
                                            Condition 1: Trademark Name Comparison\n 
                                            - Condition 1A: The existing trademark's name is a character-for-character match with the proposed trademark name.\n 
                                            - Condition 1B: The existing trademark's name is semantically equivalent to the proposed trademark name.\n 
                                            - Condition 1C: The existing trademark's name is phonetically equivalent to the proposed trademark name.\n 
                                            - Condition 1D: If both the existing trademark's name and the proposed trademark name consist of multiple words, then the first two or more words of the existing trademark's name must be phonetically equivalent to the proposed trademark name.\n
                                            - Condition 1E: Primary Position Requirement- In the context of trademark conflicts, the primary position of a trademark refers to the first word or phrase element in a multi-word or phrase trademark. For a conflict to exist between an existing trademark and a proposed trademark based on Condition 1E, the proposed trademark name must be in the primary position of the existing trademark. This means that the proposed trademark name should be the first word of the existing trademark.\n
                                                            Example:\n Existing Trademark: "STORIES AND JOURNEYS"\n Proposed Trademark: "JOURNEY"\n Analysis:\n The existing trademark "STORIES AND JOURNEYS" consists of multiple words/phrases.\n For the proposed trademark "JOURNEY" to be in conflict under Condition 1E, it must appear as the first word/phrase in the existing trademark.\n In this case, the first word/phrase in "STORIES AND JOURNEYS" is "STORIES", not "JOURNEY".\n Therefore, "JOURNEY" does not meet Condition 1E because it is not in the primary position of the existing trademark.\n
                                                            Example:\n Existing Trademark: "JOURNEY BY COMPANION"\n Proposed Trademark: "JOURNEY"\n Analysis:\n The existing trademark "JOURNEY BY COMPANION" consists of multiple words/phrases.\n For the proposed trademark "JOURNEY" to be in conflict under Condition 1E, it must appear as the first word/phrase in the existing trademark.\n In this case, the first word/phrase in "JOURNEY BY COMPANION" is "JOURNEY".\n Therefore, "JOURNEY" does meet Condition 1E because it is in the primary position of the existing trademark.\n
                                            
                                            Condition 2: Goods/Services Classification\n 
                                            - Condition 2: The existing trademark's goods/services are in the same class as those of the proposed trademark.\n
                                            
                                            Condition 3: Target Market and Products\n 
                                            - Condition 3A: The existing trademark's goods/services target the exact same products as the proposed trademark.\n 
                                            - Condition 3B: The existing trademark's goods/services target an exact market as the proposed trademark.\n
                                            
                                            If existing trademark in user given input satisfies:\n\n
                                            - Special case: If existing Trademark Status is Cancelled or Abandoned, they will automatically be considered as conflict grade low but still give the reasoning for the potential conflicts.\n\n
                                            - If the existing trademark satisfies Condition 1A, 1B, 1C, or 1D, and also satisfies the revised Condition 1E (when applicable), along with Condition 2, and both Condition 3A (Needed to be fully satisfied) and 3B (Needed to be fully satisfied), then the conflict grade should be High.\n
                                            - If the existing trademark satisfies any two of the following: Condition 1A, 1B, 1C, or 1D (with the revised Condition 1E being a necessary component for these to be considered satisfied when applicable), Condition 2, Condition 3A (Needed to be fully satisfied) and 3B (Needed to be fully satisfied), then the conflict grade should be Moderate.\n
                                            - If the existing trademark satisfies only one (or none) of the conditions: Condition 1A, 1B, 1C, 1D and (only if the revised Condition 1E is also satisfied when applicable), Condition 2, Condition 3A and 3B, then the conflict grade should be Low.\n\n
                                            
                                            Format of the Response:\n
                                            Resoning for Conflict: Reasoning for conflict in bullet points (In reasoning, if exact same goods, services and industries: list the overlaps, you should reasoning whether the good/services are overlapping or not including classes (if same as proposed trademark or not) and see trademark name whether identical (character-for-character) matches, phonetic equivalents, if it is in primary position (first word in the phrase) or not, if it is not in primary position (first word in the phrase) of the existing trademark it is not conflicting and standard plural forms for subject goods and goods that may be related or not. Reasoning should be based on provided information. Do not provide any kind of hypothetical reasoning.)\n\n
                                            
                                            Step 0: Identifying Potential Conflicts
                                            - What is the existing trademark?
                                            - What is the status of the existing trademark? 
                                            - Who is the owner of the existing trademark?
                                            - What is the class number for the existing trademark?
                                            - What is the proposed trademark?
                                            - Who is the applicant for the proposed trademark?
                                            - What is the class number for the proposed trademark?

                                            Step 1: Condition 1A - Character-for-Character Match
                                            - Does the existing trademark match the proposed trademark character-for-character?
                                            - If not, does the existing trademark form part of the proposed trademark? Specify the matching part.

                                            Step 2: Condition 1B - Visual Similarity
                                            - Do the existing and proposed trademarks look visually similar?
                                            - If yes, describe the visual elements that contribute to this similarity.
                                            
                                            Step 3: Condition 1C - Sound-Alike Match
                                            - Do the existing and proposed trademarks sound alike when spoken?
                                            - If yes, describe the similarities in their pronunciation.

                                            Step 4: Condition 1D - Sound-Alike Match for First Two or More Words
                                            - Do the existing and proposed trademarks consist of multiple words?
                                            - Do the first two or more words of the existing trademark sound alike when spoken compared to the proposed trademark?
                                            - If yes, describe the similarities in their pronunciation.
                                            
                                            Step 5: Condition 1E - Primary Position
                                            - Is the existing trademark in the primary position of the proposed trademark?
                                            - Is the primary position the beginning of the proposed trademark?
                                            - Does the proposed trademark exactly match the existing trademark in its primary position?
                                            - If the proposed trademark name is a single word or phrase, it must be fully presented in the primary position of the existing trademark for this condition to be applicable.
                                            - As if the existing trademark name is a single word, then condition 1E is not applicable.

                                            Step 6: Condition 2 - Class Overlap
                                            - Do the class numbers of the existing and proposed trademarks match?
                                            - If the class numbers are different, are the classes related in a way that might cause consumer confusion?

                                            Step 7: Condition 3A - Goods/Services Overlap
                                            - What goods/services and products are covered by the existing trademark?
                                            - What goods/services and products are covered by the proposed trademark?
                                            - Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?

                                            Step 8: Condition 3B - Target Market
                                            - Who is the target market for the goods/services covered by the existing trademark?
                                            - Who is the target market for the goods/services covered by the proposed trademark?
                                            - Is there an exact match or exact overlap in the target market for the existing and proposed trademarks?

                                            Step 9: Conflict Grade Assessment
                                            - Based on the analysis, how would you grade the potential conflict?
                                            - Provide a brief reasoning for the conflict grade.
                                            
                                            Example Analysis Using the Steps : 
                                            - Trademark Name: MH
                                            - Trademark Status: REGISTERED
                                            - Trademark Owner: ZHAO
                                            - Trademark Class Number: 3
                                            - Proposed Trademark: MH BY MOTHERHOOD
                                            - Applicant: ABC Company
                                            - Proposed Trademark Class Number: 3

                                            Step 1: Condition 1A - Character-for-Character Match
                                            - Does the existing trademark match the proposed trademark character-for-character?
                                            - No, "MH" is part of "MH BY MOTHERHOOD" but not an exact match.

                                            Step 2: Condition 1B - Visual Similarity
                                            - Do the existing and proposed trademarks look visually similar?
                                            - No, "MH" and "MH BY MOTHERHOOD" do not look visually similar.

                                            Step 3: Condition 1C - Sound-Alike Match
                                            - Do the existing and proposed trademarks sound alike when spoken?
                                            - No, The existing trademark "MH" and the proposed trademark "MH BY MOTHERHOOD" are not phonetically equivalent, although they share the term "MH."
                                            - If the existing trademark name is multi-word, then the primary word of the existing trademark name followed by a color or the name of a country, city, or number (either numeric or letter) should be considered as Condition satisfied.

                                            Step 4: Condition 1D - Sound-Alike Match for First Two or More Words
                                            - Do the existing and proposed trademarks consist of multiple words?
                                            - No, The existing trademark "MH" is a single word.
                                            - Do the first two or more words of the existing trademark sound alike when spoken compared to the proposed trademark?
                                            - No, The existing trademark "MH" is a single word, As if the existing or proposed trademarks name is a single word, then condition 1D is not applicable.
                                            
                                            Step 5: Condition 1E - Primary Position
                                            - Is the existing trademark in the primary position of the proposed trademark?
                                            - Yes, the proposed trademark "MH BY MOTHERHOOD", which has "MH" as in the primary position of the existing trademark "MH".
                                            - If the proposed trademark name is a single word or phrase, it must be fully presented in the primary position of the existing trademark for this condition to be applicable.
                                            - As if the existing trademark name is a single word, then condition 1E is not applicable.
                                            
                                            Step 6: Condition 2 - Class Overlap
                                            - Do the class numbers of the existing and proposed trademarks match?
                                            - Yes, both are in Class 3.

                                            Step 7: Condition 3A - Goods/Services Overlap
                                            - What goods/services are covered by the existing trademark?
                                            - Non-medicated cosmetic and hair care preparations.
                                            - What goods/services are covered by the proposed trademark?
                                            - Hair care preparations.
                                            - Is there an overlap between the goods/services and product of the existing and proposed trademarks?
                                            - Yes, both include exact products such as hair care preparations.

                                            Step 8: Condition 3B - Target Market
                                            - Who is the target market for the goods/services covered by the existing trademark?
                                            - Consumers interested in cosmetic and hair care products.
                                            - Who is the target market for the goods/services covered by the proposed trademark?
                                            - Consumers interested in hair care products.
                                            - Is there an overlap in the target market for the existing and proposed trademarks?
                                            - Yes, both target the exact same consumers interested in hair care products.

                                            Step 9: Conflict Grade Assessment
                                            - Based on the analysis, how would you grade the potential conflict?
                                            - High.

                                            Provide a brief reasoning for the conflict grade.
                                            - Example : Condition 1A is not satisfied since there is no character-for-character match, Conditions 1B, and 1C are not satisfied either and Condition 1D and 1E are not applicable. However Condition 2, 3A and 3B is been satisfied and also there is a class overlap and goods/services overlap, as well as a shared target market, indicating a Moderate potential for consumer confusion.
                                                                                    
                                            Conflict Grade: Based on above reasoning (Low or Moderate or High)."""
                                            },
            
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: DISCOVER WHAT’S NEXT\n
                                            Goods/Services: Online retail store services geared for men and women, featuring a wide variety of unique consumer products\n 
                                            International Class Numbers: 35\n
                                            Status: REGISTERED\n
                                            Owner: THE GIDDYUP GROUP, INC\n
                                            
                                            Proposed Trademark:\n
                                            Name: DISCOVER\n 
                                            Goods/Services: Luggage and carrying bags; suitcases, trunks, travelling bags, sling bags for = carrying infants, school bags; purses; wallets; retail and online retail services\n
                                            International Class Numbers: 18, 35\n"""
            },
        {"role": "assistant", "content":"""
Reasoning for Conflict:
Step 1: Condition 1A - Exact Character Match
- Does the existing trademark contain the proposed trademark term?
- Yes, the existing trademark "DISCOVER WHAT’S NEXT" contains the term "DISCOVER," which is a character for-character match with the proposed trademark "DISCOVER."
- Condition 1A is satisfied.

Step 2: Condition 1E - Primary Position
- Is the existing trademark in the primary position of the proposed trademark?
- Yes, "DISCOVER" is in the primary position of the existing trademark "DISCOVER WHAT’S NEXT."
- As if the existing trademark name is a single word, then condition 1E is not applicable.
- Existing trademark name is a single word, Hence Condition 1E is not applicable.

Step 3: Condition 2 - Class Overlap
- Do the class numbers of the existing and proposed trademarks overlap?
- Yes, the existing trademark is registered under International Class 35, which overlaps with the proposed trademark's Class 35 for retail and online retail services.
- Condition 2 is satisfied.

Step 4: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services include online retail store services geared for men and women, featuring a wide variety of unique consumer products.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services include luggage and carrying bags; suitcases, trunks, travelling bags, sling bags for carrying infants, school bags; purses; wallets; retail and online retail services.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- No, while both trademarks include retail and online retail services, the existing trademark focuses on a wide variety of unique consumer products, whereas the proposed trademark focuses specifically on luggage, bags, and related items.
- Condition 3A is not fully satisfied.

Step 5: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in a wide variety of unique consumer products.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in luggage, carrying bags, and related products.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- No, although there is some overlap in the general market of retail and online retail services, the specific focus of the products differs.
- Condition 3B is not fully satisfied.

Conflict Reason:
Reasoning for Conflict:
1A - The existing trademark "DISCOVER WHAT’S NEXT" contains the term "DISCOVER," which is a character for-character match with the proposed trademark "DISCOVER."
1E - "DISCOVER" is in the primary position of the existing trademark "DISCOVER WHAT’S NEXT". However, The existing trademark name is a single word, Hence Condition 1E is not applicable. 
2  - The existing trademark is registered under International Class 35, which overlaps with the proposed trademark's Class 35 for retail and online retail services.
3A - Both trademarks include retail and online retail services, the existing trademark focuses on a wide variety of unique consumer products, whereas the proposed trademark focuses specifically on luggage, bags, and related items.
3B - Although there is some overlap in the general market of retail and online retail services, the specific focus of the products differs.

Conclusion:
- Despite satisfying Conditions 1A, (1E not applicable) , and 2, the differences in the specific goods/services and target markets mean that the conflict is not high.
- Since Conditions 3A and 3B are not fully satisfied due to the lack of exact overlap in the goods/services and target markets, the conflict grade should be moderate.

- Conflict Grade: Moderate

"""
            },
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: DB JOURNEY\n
                                            Goods/Services: All-purpose carrying bags, rucksacks, hipsacks, luggage, toiletry bags, key bags, luggage tags, pocket wallets, straps for luggage, shoulder straps, and umbrellas (Class 18)\n 
                                            International Class Numbers: 9, 16, 18, 25, 28\n
                                            Status: PENDING SECTION 66(A) (MADRID PROTOCOL)\n
                                            Owner: DB EQUIPMENT AS NORWAY AS\n
                                            
                                            Proposed Trademark:\n
                                            Name: JOURNEY\n 
                                            Goods/Services: Luggage and carrying bags; suitcases, trunks, travelling bags, sling bags for carrying infants, school bags; purses; wallets; retail and online retail services\n
                                            International Class Numbers: 18, 35\n"""
            },
            {"role": "assistant", "content":""" 
Reasoning for Conflict:
Step 1: Condition 1A - Character-for-Character Match
- The existing trademark "DB JOURNEY" is not a character-for-character match with the proposed trademark "JOURNEY."
- Condition 1A is not satisfied.

Step 2: Condition 1B - Semantic Equivalence
- The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are not semantically equivalent.
- Condition 1B is not satisfied.

Step 3: Condition 1C - Phonetic Equivalence
- The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are not phonetically equivalent, although they share the term "JOURNEY."
- Condition 1C is not satisfied.

Step 4: Condition 1D - Phonetic Equivalence for First Two or More Words
- The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are not phonetically equivalent, although they share the term "JOURNEY."
- The proposed trademark "JOURNEY" is a single word, As if the existing or proposed trademarks name is a single word, then condition 1D is not applicable.
- Condition 1D is not applicable.

Step 5: Condition 1E - Primary Position
- The term "JOURNEY" is in the primary position in the proposed trademark "JOURNEY."
- The term "JOURNEY" is not in the primary position in the existing trademark "DB JOURNEY" (the primary term is "DB").
- Condition 1E is not satisfied.

Step 6: Condition 2 - Class Overlap
- The existing trademark includes Class 18, which overlaps with the proposed trademark's Class 18 for luggage and carrying bags.
- The existing trademark does not include Class 35, which is part of the proposed trademark's classification.
- Condition 2 is satisfied.

Step 7: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services in Class 18 include all-purpose carrying bags, rucksacks, hipsacks, luggage, toiletry bags, key bags, luggage tags, pocket wallets, straps for luggage, shoulder straps, and umbrellas.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services in Class 18 include luggage and carrying bags, suitcases, trunks, travelling bags, sling bags for carrying infants, school bags, purses, and wallets.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- There is an exact overlap in the goods/services in Class 18.
- Condition 3A is satisfied.

Step 8: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in a wide range of bags and carrying cases, including luggage and related accessories.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in luggage, carrying bags, and related products.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- There is an overlap in the target market for the existing and proposed trademarks.
- Condition 3B is satisfied.

Conflict Reason:
Reasoning for Conflict:
1A - The existing trademark "DB JOURNEY" is not a character-for-character match with the proposed trademark "JOURNEY."
1B - The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are not semantically equivalent.
1C - The existing trademark "DB JOURNEY" and the proposed trademark "JOURNEY" are not phonetically equivalent, although they share the term "JOURNEY."
1D - The proposed trademark "JOURNEY" is a single word, As if the existing or proposed trademarks name is a single word, then condition 1D is not applicable.
1E - The term "JOURNEY" is not in the primary position in the existing trademark "DB JOURNEY" (the primary term is "DB").
2  - The existing trademark includes Class 18, which overlaps with the proposed trademark's Class 18 for luggage and carrying bags.
3A - There is an exact overlap in the goods/services in Class 18.
3B - There is an overlap in the target market for the existing and proposed trademarks.

Conclusion:
- Conditions 2, 3A, and 3B are satisfied.
- Condition 1A, 1B, 1C and 1E is not satisfied due to the proposed trademark name "JOURNEY" not being in the primary position of the existing trademark name "DB JOURNEY" and Condition 1D is not applicable.
- Given the significant overlap in goods/services and the shared target market but the non-satisfaction of Condition 1A or 1B or 1C or 1D with 1E, the conflict grade should be moderate.

- Conflict Grade: Moderate

"""
            },    
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: SCOOPT'D\n
                                            Goods/Services: Hypoallergenic and vegan-friendly ice cream\n 
                                            International Class Numbers: 30\n
                                            Status: Registered\n
                                            Owner: Scoopt'd: Dairy Free Treats\n
                                            
                                            Proposed Trademark:\n
                                            Name: SCOOP-A-PALOOZA\n 
                                            Goods/Services: Ice cream\n
                                            International Class Numbers: 30\n"""
            },
            {"role": "assistant", "content":""" 
Reasoning for Conflict:
Step 1: Condition 1A - Character-for-Character Match
- The existing trademark "SCOOPT'D" is not a character-for-character match with the proposed trademark "SCOOP-A-PALOOZA."
- Condition 1A is not satisfied.

Step 2: Condition 1B - Semantic Equivalence
- The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are not semantically equivalent.
- Condition 1B is not satisfied.

Step 3: Condition 1C - Phonetic Equivalence
- The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are partially phonetically similar due to the shared term "SCOOP," which could lead to confusion. However, they are not fully phonetically similar.
- Condition 1C is not satisfied.

Step 4: Condition 1D - Phonetic Equivalence for First Two or More Words
- The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are partially phonetically similar due to the shared term "SCOOP," which could lead to confusion. However, they are not fully phonetically similar.
- Both the existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are single words, As if the existing or proposed trademarks name is a single word, then condition 1D is not applicable.
- Condition 1D is not applicable.

Step 5: Condition 1E - Primary Position
- The term "SCOOP" is in the primary position in the proposed trademark "SCOOP-A-PALOOZA."
- The term "SCOOP" is in the primary position in the existing trademark "SCOOPT'D."
- The full proposed trademark "SCOOP-A-PALOOZA" does not match the primary position of the existing trademark "SCOOPT'D."
- As if the existing trademark name is a single word, then condition 1E is not applicable.
- Existing trademark name is a single word, Hence Condition 1E is not applicable.
- Condition 1E is not applicable.

Step 6: Condition 2 - Class Overlap
- Both the existing and proposed trademarks are in International Class 30.
- Condition 2 is satisfied.

Step 7: Condition 3A - Goods/Services Overlap
- What goods/services and products are covered by the existing trademark?
- The existing trademark's goods/services include hypoallergenic and vegan-friendly ice cream.
- What goods/services and products are covered by the proposed trademark?
- The proposed trademark's goods/services include ice cream.
- Is there an exact match or exact overlap between the goods/services and products of the existing and proposed trademarks?
- Yes, both trademarks cover ice cream.
- Condition 3A is satisfied.

Step 8: Condition 3B - Target Market Overlap
- What is the target market for the existing trademark?
- The existing trademark targets consumers interested in hypoallergenic and vegan-friendly ice cream.
- What is the target market for the proposed trademark?
- The proposed trademark targets consumers interested in ice cream.
- Is there an exact overlap in the target market for the existing and proposed trademarks?
- Yes, both trademarks target consumers interested in ice cream.
- Condition 3B is satisfied.

Conflict Reason:
Reasoning for Conflict:
1A - The existing trademark "SCOOPT'D" is not a character-for-character match with the proposed trademark "SCOOP-A-PALOOZA."
1B - The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are not semantically equivalent.
1C - The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are partially phonetically similar due to the shared term "SCOOP," which could lead to confusion. However, they are not fully phonetically similar.
1D - Both the existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" are single words, As if the existing or proposed trademarks name is a single word, then condition 1D is not applicable.
1E - The full proposed trademark "SCOOP-A-PALOOZA" does not match the primary position of the existing trademark "SCOOPT'D." However, Existing trademark name is a single word, Hence Condition 1E is not applicable.
2  - Both the existing and proposed trademarks are in International Class 30.
3A - Both trademarks cover ice cream.
3B - Both trademarks target consumers interested in ice cream.

Conclusion:
- Conditions 2, 3A, and 3B and (1D and 1E not applicable) are satisfied.
- Conditions 1A, 1B, 1C are not satisfied.
- Given the overlap in goods/services and the shared target market but the lack of satisfaction in key similarity conditions (1A, 1B, 1C, 1D, and 1E), the conflict grade should be moderate.

- Conflict Grade: Moderate
"""
            },
            {"role": "user", "content": f"""Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: {existing_trademark['trademark_name']}\n
                                            Goods/Services: {existing_trademark['goods_services']}\n 
                                            International Class Numbers: {existing_trademark['international_class_number']}\n
                                            Status: {existing_trademark['status']}\n
                                            Owner: {existing_trademark['owner']}\n
                                            Proposed Trademark:\n
                                            Name: {proposed_name}\n 
                                            Goods/Services: {proposed_goods_services}\n
                                            International Class Numbers: {proposed_classes}\n"""
            }
        ]

    from openai import AzureOpenAI
    client = AzureOpenAI(  
                azure_endpoint="https://danielingitaraj.openai.azure.com/",  
                api_key="a5c4e09a50dd4e13a69e7ef19d07b48c",  
                api_version="2024-02-01",
            )  
                
    response_reasoning = client.chat.completions.create(  
                        model="GPT4",  
                        messages=messages,  
                        temperature=0,  
                        max_tokens=4095,  
                        top_p = 1
                    )

    Treasoning = response_reasoning.choices[0].message.content
    print(Treasoning)
    print("_____________________________________________________________________________________________________________________________")
    reasoning = Treasoning.split("Conflict Reason:", 1)[1].strip()
    conflict_grade = Treasoning.split("Conflict Grade:", 1)[1].strip() 
    progress_bar.progress(60)
    
    return {
        'Trademark name': existing_trademark['trademark_name'],
        'Trademark status': existing_trademark['status'],
        'Trademark owner': existing_trademark['owner'],
        'Trademark class Number': existing_trademark['international_class_number'],
        'Trademark serial number' : existing_trademark['serial_number'],
        'Trademark registration number' : existing_trademark['registration_number'],
        'Trademark design phrase' : existing_trademark['design_phrase'],
        'conflict_grade': conflict_grade,
        'reasoning': reasoning
    }



def compare_trademarks2(existing_trademark: List[Dict[str, Union[str, List[int]]]], proposed_name: str, proposed_class: str, proposed_goods_services: str) -> List[Dict[str, int]]:
    proposed_classes = [int(c.strip()) for c in proposed_class.split(',')]
    messages=[
            {"role": "system", "content": """You are a trademark attorney tasked with determining a conflict grade based on the given conditions. You should assign a conflict grade of "Name-Match" or "Low" to the existing trademark and respond with only "Name-Match", or "Low".
                                            Conditions for Determining Conflict Grades:

                                            Condition 1: Trademark Name Comparison
                                            - Condition 1A: Are the existing trademark name and the proposed trademark name in conflict with respect to Distinctiveness, Strength of the Marks, and Similarity in Appearance, Sound, and Meaning?

                                            If the existing trademark in the user-provided input satisfies:
                                            - Special Case: If the existing trademark status is "Cancelled" or "Abandoned," it will automatically be considered a conflict grade of "Low," but you should still provide reasoning for any potential conflicts.
                                            - If the existing trademark satisfies Condition 1A, then the conflict grade should be "Name-Match."
                                            - If the existing trademark does not satisfy Condition 1A, then the conflict grade should be "Low."

                                            Format of the Response:
                                            Reasoning for Conflict: Provide reasoning for the conflict in bullet points. In your reasoning, if the goods, services, and industries are exactly the same, list the overlaps. You should determine whether the goods/services overlap, including classes (whether they are the same as the proposed trademark or not). Consider whether the trademark names are identical (character-for-character matches), phonetic equivalents, if the name is in the primary position (first word in the phrase), or if it is not in the primary position of the existing trademark. If it is not in the primary position, it is not conflicting. Also, consider standard plural forms for subject goods and whether the goods may be related or not. Reasoning should be based on the provided information. Do not provide any hypothetical reasoning.
                                            Note: Also mention if the existing trademark and the proposed trademark are not in the same Class number in the Reasoning for Conflict.
                                            
                                            Step 0: Identifying Potential Conflicts
                                            - What is the existing trademark?
                                            - What is the status of the existing trademark? 
                                            - What is the proposed trademark?

                                            Step 1: Check the Status of the Existing Trademark:
                                            - If the existing trademark is "Cancelled" or "Expired" or "Abandoned," assign the conflict grade as "Low." And skip other conditions.

                                            Step 2: **Trademark Name Comparison:**
                                            - Evaluate if there is a conflict between the existing trademark name and the proposed trademark name based on the following:
                                                - **Distinctiveness and Strength of the Marks:** Are the trademarks distinctive or similar in strength?
                                                - **Similarity in Appearance, Sound, and Meaning:** Do the names look, sound, or mean the same?

                                            Step 3: **Consider Special Cases and Additional Factors:**
                                            - Consider standard plural forms for the subject goods and whether goods may be related or not.
                                            - If there is no similarity in name, class, or overlapping goods/services, assign the conflict grade "Low."

                                            Format of the Response:
                                            - **Reasoning for Conflict:** Provide reasoning in bullet points. Base your reasoning only on the provided information. Clearly mention if the existing and proposed trademarks are not in the same Class number.

                                            Note:
                                            - Do not provide any hypothetical reasoning. The conflict grade should be based solely on the facts given.
                                            
                                            Example Analysis Using the Steps : 
                                            - Trademark Name: Unlock Brisk's Bold Flavors
                                            - Trademark Status: REGISTERED
                                            - Proposed Trademark: Unlock Hidden Flavors

                                            Reasoning for Conflict:
                                            Step 1: Status Check: 
                                            - The existing trademark status is not "Cancelled" or "Expired" or "Abandoned,". Proceeding to name comparison.

                                            Step 2: Trademark Name Comparison:
                                            - Both trademarks share the distinctive word "Unlock" in the primary position, which creates a similarity in appearance, sound, and meaning.
                                            - The phrase "Unlock Flavors" forms the core part of both trademarks, creating a strong conceptual similarity.
                                            - While the words "Brisk's Bold" and "Hidden" differ, they serve as modifiers to the common term "Flavors." The similarity is substantial because the focus of both trademarks is on the concept of "Unlocking Flavors," which could confuse consumers regarding the source or affiliation of the products/services.

                                            Step 3: No special cases apply (such as status being "Cancelled" or "Abandoned").
                                            
                                            Conflict Reason:
                                            Reasoning for Conflict:
                                            The shared use of the distinctive phrase "Unlock Flavors" as the primary conceptual focus in both trademarks creates a strong similarity in appearance, sound, and meaning. This overlap is significant enough to potentially confuse consumers about the source or affiliation of the goods or services, thereby assigning the conflict grade as "Name-Match."
                                            
                                            - Conflict Grade: Name-Match
                                            """
                                            },            
            {"role": "user", "content": """Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: SCOOPT'D\n
                                            Status: Registered\n
                                                                                        
                                            Proposed Trademark:\n
                                            Name: SCOOP-A-PALOOZA\n """
            },
            {"role": "assistant", "content":""" 
Reasoning for Conflict:
Step 1: Status Check:
- The existing trademark status is not "Cancelled," "Expired," or "Abandoned." Proceeding to name comparison.

Step 2: Trademark Name Comparison:
- Both trademarks share the distinctive word "SCOOP," which creates a similarity in appearance, sound, and meaning.
- The word "SCOOP" is the dominant and distinctive part of both trademarks, leading to a conceptual similarity.
- The existing trademark "SCOOPT'D" and the proposed trademark "SCOOP-A-PALOOZA" both emphasize the idea of "Scoop," likely related to ice cream or a similar product, which could confuse consumers regarding the source or affiliation.
- The additional elements ("T'D" and "A-PALOOZA") differ but serve as suffixes or modifiers to the common term "SCOOP."

Step 3: No special cases apply (such as status being "Cancelled" or "Abandoned").

Conflict Reason:
Reasoning for Conflict:
The shared use of the distinctive word "SCOOP" as the core focus in both trademarks creates a strong similarity in appearance, sound, and meaning. This overlap is significant enough to potentially confuse consumers about the source or affiliation of the goods or services, thereby assigning the conflict grade as "Name-Match."

- Conflict Grade: Name-Match
"""
            },
            {"role": "user", "content": f"""Compare the following existing and proposed trademarks and determine the conflict grade.\n
                                            Existing Trademark:\n
                                            Name: {existing_trademark['trademark_name']}\n
                                            Status: {existing_trademark['status']}\n
                                            
                                            Proposed Trademark:\n
                                            Name: {proposed_name}\n """
            }
        ]

    from openai import AzureOpenAI
    client = AzureOpenAI(  
                azure_endpoint="https://danielingitaraj.openai.azure.com/",  
                api_key="a5c4e09a50dd4e13a69e7ef19d07b48c",  
                api_version="2024-02-01",
            )  
                
    response_reasoning = client.chat.completions.create(  
                        model="GPT4",  
                        messages=messages,  
                        temperature=0,  
                        max_tokens=4095,  
                        top_p = 1
                    )

    Treasoning = response_reasoning.choices[0].message.content
    print(Treasoning)
    print("_____________________________________________________________________________________________________________________________")
    reasoning = Treasoning.split("Reasoning for Conflict:", 1)[1].strip()
    conflict_grade = Treasoning.split("Conflict Grade:", 1)[1].strip() 
    
    return {
        'Trademark name': existing_trademark['trademark_name'],
        'Trademark status': existing_trademark['status'],
        'Trademark owner': existing_trademark['owner'],
        'Trademark class Number': existing_trademark['international_class_number'],
        'Trademark serial number' : existing_trademark['serial_number'],
        'Trademark registration number' : existing_trademark['registration_number'],
        'Trademark design phrase' : existing_trademark['design_phrase'],
        'conflict_grade': conflict_grade,
        'reasoning': reasoning
    }
    

def extract_proposed_trademark_details(file_path: str) -> Dict[str, Union[str, List[int]]]:
    """ Extract proposed trademark details from the given input format """
    proposed_details = {}
    with fitz.open(file_path) as pdf_document:
        if pdf_document.page_count > 0:
            page = pdf_document.load_page(0)
            page_text = preprocess_text(page.get_text())
            
    name_match = re.search(r'Mark Searched:\s*(.*?)(?=\s*Client Name:)', page_text, re.IGNORECASE | re.DOTALL)
    if name_match:
        proposed_details["proposed_trademark_name"] = name_match.group(1).strip()

    goods_services_match = re.search(r'Goods/Services:\s*(.*?)(?=\s*Trademark Research Report)', page_text, re.IGNORECASE | re.DOTALL)
    if goods_services_match:
        proposed_details["proposed_goods_services"] = goods_services_match.group(1).strip()
    
    # Use LLM to find the international class number based on goods & services
    if "proposed_goods_services" in proposed_details:
        goods_services = proposed_details["proposed_goods_services"]
        class_numbers = find_class_numbers(goods_services)
        proposed_details["proposed_nice_classes_number"] = class_numbers
    
    return proposed_details

def find_class_numbers(goods_services: str) -> List[int]:
    """ Use LLM to find the international class numbers based on goods & services """
        # Initialize AzureChatOpenAI
    
    from openai import AzureOpenAI
    client = AzureOpenAI(  
                    azure_endpoint="https://theswedes.openai.azure.com/",  
                    api_key="783973291a7c4a74a1120133309860c0",  
                    api_version="2024-02-01",
    )  
    messages=[
            {"role": "system", "content": "You are a helpful assistant for finding the International class number of provided Goods & Services."},
            {"role": "user", "content": "The goods/services are: IC 003: SKIN CARE PREPARATIONS; COSMETICS; BABY CARE PRODUCTS, NAMELY, SKIN SOAPS, BABY WASH, BABY BUBBLE BATH, BABY LOTIONS, BABY SHAMPOOS; SKIN CLEANSERS; BABY WIPES; NON− MEDICATED DIAPER RASH OINTMENTS AND LOTIONS; SKIN LOTIONS, CREAMS, MOISTURIZERS, AND OILS; BODY WASH; BODY SOAP; DEODORANTS; PERFUME; HAIR CARE PREPARATIONS. Find the international class numbers."},
            {"role": "assistant", "content": "The international class numbers : 03"},
            {"role": "user", "content": "The goods/services are: LUGGAGE AND CARRYING BAGS; SUITCASES, TRUNKS, TRAVELLING BAGS, SLING BAGS FOR CARRYING INFANTS, SCHOOL BAGS; PURSES; WALLETS; RETAIL AND ONLINE RETAIL SERVICES. Find the international class numbers."},
            {"role": "assistant", "content": "The international class numbers : 18,35"},
            {"role": "user", "content": f"The goods/services are: {goods_services}. Find the international class numbers."}
    ]
    response = client.chat.completions.create(  
                        model="GPT35",  
                        messages=messages,  
                        temperature=0,  
                        max_tokens=150,  
    )  
    
    class_numbers_str = response.choices[0].message.content
    
    # Extracting class numbers and removing duplicates
    class_numbers = re.findall(r'(?<!\d)\d{2}(?!\d)', class_numbers_str)  # Look for two-digit numbers
    class_numbers = ','.join(set(class_numbers))  # Convert to set to remove duplicates, then join into a single string
    
    return class_numbers

def extract_proposed_trademark_details2(file_path: str) -> Dict[str, Union[str, List[int]]]:
    """ Extract proposed trademark details from the first page of the document """
    proposed_details = {}
    with fitz.open(file_path) as pdf_document:
        if pdf_document.page_count > 0:
            page = pdf_document.load_page(0)
            page_text = preprocess_text(page.get_text())
            
            name_match = re.search(r'Name:\s*(.*?)(?=\s*Nice Classes:)', page_text)
            if name_match:
                proposed_details["proposed_trademark_name"] = name_match.group(1).strip()
                
            nice_classes_match = re.search(r'Nice Classes:\s*(\d+(?:,\s*\d+)*)', page_text)
            if nice_classes_match:
                proposed_details["proposed_nice_classes_number"] = nice_classes_match.group(1).strip()
            
            goods_services_match = re.search(r'Goods & Services:\s*(.*?)(?=\s*Registers|$)', page_text, re.IGNORECASE | re.DOTALL)
            if goods_services_match:
                proposed_details["proposed_goods_services"] = goods_services_match.group(1).strip()
    
    return proposed_details

def list_conversion(proposed_class: str) -> List[int]:
    
    from openai import AzureOpenAI
    client = AzureOpenAI(  
                    azure_endpoint="https://theswedes.openai.azure.com/",  
                    api_key="783973291a7c4a74a1120133309860c0",  
                    api_version="2024-02-01",
    )  
    messages=[
                    {"role": "system", "content": "You are a helpful assistant for converting the class number string into python list of numbers.\n Respond only with python list. Example : [18,35]"},
                    {"role": "user", "content": "The class number are: 15,89. convert the string into python list of numbers."},
                    {"role": "assistant", "content": "[15,89]"},
                    {"role": "user", "content": f"The class number are: {proposed_class}. convert the string into python list of numbers."}
                ]
    response = client.chat.completions.create(  
                        model="GPT35",  
                        messages=messages,  
                        temperature=0,  
                        max_tokens=150,  
    )  

    lst_class = response.choices[0].message.content
    class_value = ast.literal_eval(lst_class)
            
    return class_value

# Streamlit App  
st.title("Trademark Document Parser Version 6.2")  
  
# File upload  
uploaded_files = st.sidebar.file_uploader("Choose PDF files", type="pdf", accept_multiple_files=True)  
  
if uploaded_files:  
    if st.sidebar.button("Check Conflicts", key="check_conflicts"):  
        total_files = len(uploaded_files)  
        progress_bar = st.progress(0)  
        for i, uploaded_file in enumerate(uploaded_files):  
            # Save uploaded file to a temporary file path  
            temp_file_path = f"temp_{uploaded_file.name}"  
            with open(temp_file_path, "wb") as f:  
                f.write(uploaded_file.read())  
            
            sp = True
            proposed_trademark_details = extract_proposed_trademark_details(temp_file_path)  
                            
            if proposed_trademark_details:  
                proposed_name = proposed_trademark_details.get('proposed_trademark_name', 'N')  
                proposed_class = proposed_trademark_details.get('proposed_nice_classes_number')  
                proposed_goods_services = proposed_trademark_details.get('proposed_goods_services', 'N') 
                if (proposed_goods_services != 'N'): 
                    with st.expander(f"Proposed Trademark Details for {uploaded_file.name}"):  
                            st.write(f"Proposed Trademark name: {proposed_name}")  
                            st.write(f"Proposed class-number: {proposed_class}")  
                            st.write(f"Proposed Goods & Services: {proposed_goods_services}") 
                    class_list = list_conversion(proposed_class) 
                else :
                    st.write("______________________________________________________________________________________________________________________________")
                    st.write(f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}")
                    st.write("______________________________________________________________________________________________________________________________")
                    sp = False
            else:  
                
                proposed_trademark_details = extract_proposed_trademark_details2(temp_file_path)  
                
                if proposed_trademark_details:  
                    proposed_name = proposed_trademark_details.get('proposed_trademark_name', 'N')  
                    proposed_class = proposed_trademark_details.get('proposed_nice_classes_number')  
                    proposed_goods_services = proposed_trademark_details.get('proposed_goods_services', 'N')  
                    if (proposed_goods_services != 'N'): 
                        with st.expander(f"Proposed Trademark Details for {uploaded_file.name}"):  
                                st.write(f"Proposed Trademark name: {proposed_name}")  
                                st.write(f"Proposed class-number: {proposed_class}")  
                                st.write(f"Proposed Goods & Services: {proposed_goods_services}") 
                        class_list = list_conversion(proposed_class)  
                    else :
                        st.write("______________________________________________________________________________________________________________________________")
                        st.write(f"Sorry, unable to generate report due to insufficient information about goods & services in the original trademark report : {uploaded_file.name}")
                        st.write("______________________________________________________________________________________________________________________________")
                        sp = False
                else :  
                    st.error(f"Unable to extract Proposed Trademark Details for {uploaded_file.name}") 
                    sp = False 
                    continue  
            
            if (sp):    
                for i in range(1,21):
                    time.sleep(0.5)
                    progress_bar.progress(i)
                    
                progress_bar.progress(25)
                # Initialize AzureChatOpenAI
                
                existing_trademarks = parse_trademark_details(temp_file_path)
                for i in range(25,46):
                    time.sleep(0.5)
                    progress_bar.progress(i)  
                    
                progress_bar.progress(50)
                st.success(f"Existing Trademarks Data Extracted Successfully for {uploaded_file.name}!")  
                # Display extracted details              
                
                nfiltered_list = []
                unsame_class_list = []
                
                # Iterate over each JSON element in trademark_name_list  
                for json_element in existing_trademarks:  
                    class_numbers = json_element["international_class_number"]  
                # Check if any of the class numbers are in class_list  
                    if any(number in class_list for number in class_numbers):  
                        nfiltered_list.append(json_element)
                    else:
                        unsame_class_list.append(json_element)
                    
                existing_trademarks = nfiltered_list  
                existing_trademarks_unsame =  unsame_class_list
                     
                high_conflicts = []
                moderate_conflicts = []
                low_conflicts = []
                Name_Matchs = []
                
                lt = len(existing_trademarks)
                
                for existing_trademark in existing_trademarks:  
                    conflict = compare_trademarks(existing_trademark, proposed_name, proposed_class, proposed_goods_services)  
                    if conflict['conflict_grade'] == "High":  
                        high_conflicts.append(conflict)  
                    elif conflict['conflict_grade'] == "Moderate":  
                        moderate_conflicts.append(conflict)  
                    else:  
                        low_conflicts.append(conflict)  
                        
                for existing_trademarks in existing_trademarks_unsame:
                    conflict = compare_trademarks2(existing_trademarks, proposed_name, proposed_class, proposed_goods_services)  

                    if conflict['conflict_grade'] == "Name-Match":  
                        Name_Matchs.append(conflict)  
                    else:  
                        print("Low")
                        # low_conflicts.append(conflict) 
                        
                st.sidebar.write("_________________________________________________")
                st.sidebar.subheader("\n\nConflict Grades : \n")  
                st.sidebar.markdown(f"File: {proposed_name}")  
                st.sidebar.markdown(f"Total number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(Name_Matchs) + len(low_conflicts)}")
                st.sidebar.markdown(f"High Conflicts: {len(high_conflicts)}")  
                st.sidebar.markdown(f"Moderate Conflicts: {len(moderate_conflicts)}")  
                st.sidebar.markdown(f"Name Match's Conflicts: {len(Name_Matchs)}")  
                st.sidebar.markdown(f"Low Conflicts: {len(low_conflicts)}")  
                st.sidebar.write("_________________________________________________")
    
                document = Document()  
                
                document.add_heading(f'Trademark Conflict List for {proposed_name} (VERSION - 6.2) :')            
                document.add_paragraph(f"\n\nTotal number of conflicts: {len(high_conflicts) + len(moderate_conflicts) + len(Name_Matchs) + len(low_conflicts)}\n- High Conflicts: {len(high_conflicts)}\n- Moderate Conflicts: {len(moderate_conflicts)}\n- Name Match's Conflicts: {len(Name_Matchs)}\n- Low Conflicts: {len(low_conflicts)}\n")  
                
                if len(high_conflicts) > 0:  
                            document.add_heading('Trademarks with High Conflicts:', level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_high = pd.DataFrame(high_conflicts) 
                            df_high = df_high.drop(columns=['Trademark serial number','Trademark registration number','Trademark design phrase','reasoning'])  
                            # Create a table in the Word document    
                            table_high = document.add_table(df_high.shape[0] + 1, df_high.shape[1])
                            # Set a predefined table style (with borders)  
                            table_high.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_high.columns):  
                                table_high.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_high.iterrows():  
                                for j, value in enumerate(row):  
                                    table_high.cell(i + 1, j).text = str(value)

                if len(moderate_conflicts) > 0:  
                            document.add_heading('Trademarks with Moderate Conflicts:', level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_moderate = pd.DataFrame(moderate_conflicts)
                            df_moderate = df_moderate.drop(columns=['Trademark serial number','Trademark registration number','Trademark design phrase','reasoning'])  
                            # Create a table in the Word document    
                            table_moderate = document.add_table(df_moderate.shape[0] + 1, df_moderate.shape[1])
                            # Set a predefined table style (with borders)  
                            table_moderate.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_moderate.columns):  
                                table_moderate.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_moderate.iterrows():  
                                for j, value in enumerate(row):  
                                    table_moderate.cell(i + 1, j).text = str(value)
                                    
                if len(Name_Matchs) > 0:  
                            document.add_heading("Trademarks with Name Match's Conflicts:", level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_Name_Matchs = pd.DataFrame(Name_Matchs)
                            df_Name_Matchs = df_Name_Matchs.drop(columns=['Trademark serial number','Trademark registration number','Trademark design phrase','reasoning'])  
                            # Create a table in the Word document    
                            table_Name_Matchs = document.add_table(df_Name_Matchs.shape[0] + 1, df_Name_Matchs.shape[1])
                            # Set a predefined table style (with borders)  
                            table_Name_Matchs.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_Name_Matchs.columns):  
                                table_Name_Matchs.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_Name_Matchs.iterrows():  
                                for j, value in enumerate(row):  
                                    table_Name_Matchs.cell(i + 1, j).text = str(value)

                if len(low_conflicts) > 0:  
                            document.add_heading('Trademarks with Low Conflicts:', level=2)  
                            # Create a pandas DataFrame from the JSON list    
                            df_low = pd.DataFrame(low_conflicts)  
                            df_low = df_low.drop(columns=['Trademark serial number','Trademark registration number','Trademark design phrase','reasoning'])
                            # Create a table in the Word document    
                            table_low = document.add_table(df_low.shape[0] + 1, df_low.shape[1])
                            # Set a predefined table style (with borders)  
                            table_low.style = 'TableGrid'  # This is a built-in style that includes borders  
                            # Add the column names to the table    
                            for i, column_name in enumerate(df_low.columns):  
                                table_low.cell(0, i).text = column_name  
                            # Add the data to the table    
                            for i, row in df_low.iterrows():  
                                for j, value in enumerate(row):  
                                    table_low.cell(i + 1, j).text = str(value)
                            
                def add_conflict_paragraph(document, conflict):  
                    p = document.add_paragraph(f"Trademark Name : {conflict.get('Trademark name', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark Status : {conflict.get('Trademark status', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark Owner : {conflict.get('Trademark owner', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark Class Number : {conflict.get('Trademark class Number', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)
                    p.paragraph_format.space_after = Pt(0)  
                    p = document.add_paragraph(f"Trademark serial number : {conflict.get('Trademark serial number', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(f"Trademark registration number : {conflict.get('Trademark registration number', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0) 
                    p = document.add_paragraph(f"Trademark Design phrase : {conflict.get('Trademark design phrase', 'N/A')}")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0)
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p.paragraph_format.space_after = Pt(0) 
                    p = document.add_paragraph(f"{conflict.get('reasoning','N/A')}\n")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                
                if len(high_conflicts) > 0:  
                    document.add_heading('Trademarks with High Conflicts Reasoning:', level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in high_conflicts:  
                        add_conflict_paragraph(document, conflict)  
                
                if len(moderate_conflicts) > 0:  
                    document.add_heading('Trademarks with Moderate Conflicts Reasoning:', level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in moderate_conflicts:  
                        add_conflict_paragraph(document, conflict)  
                        
                if len(Name_Matchs) > 0:  
                    document.add_heading("Trademarks with Name Match's Conflicts Reasoning:", level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in Name_Matchs:  
                        add_conflict_paragraph(document, conflict)  
                
                if len(low_conflicts) > 0:  
                    document.add_heading('Trademarks with Low Conflicts Reasoning:', level=2)  
                    p = document.add_paragraph(" ")  
                    p.paragraph_format.line_spacing = Pt(18)  
                    for conflict in low_conflicts:  
                        add_conflict_paragraph(document, conflict)  
                        
                for i in range(70,96):
                    time.sleep(0.5)
                    progress_bar.progress(i)  
                    
                progress_bar.progress(100)
    
                filename = proposed_name
                doc_stream = BytesIO()  
                document.save(doc_stream)  
                doc_stream.seek(0)  
                download_table = f'<a href="data:application/octet-stream;base64,{base64.b64encode(doc_stream.read()).decode()}" download="{filename + " Trademark Conflict Report"}.docx">Download: {filename}</a>'  
                st.sidebar.markdown(download_table, unsafe_allow_html=True)  
                st.success(f"{proposed_name} Document conflict report successfully completed!")
                st.write("______________________________________________________________________________________________________________________________")
  
        progress_bar.progress(100)
        st.success("All documents processed successfully!")